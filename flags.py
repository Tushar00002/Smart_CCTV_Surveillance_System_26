from docx import Document
from docx.shared import Pt
import datetime

#flags
prohibited_email = False
restricted_email = False
fire_email = False
accident_email = False
started = False
name = "_"
email = "_"

def gen_report():
    time = datetime.datetime.now().strftime("%H:%M  %a,  %b%y")
    # Load the existing .docx file
    doc = Document("reports/report.docx")
    # Define the default font size (e.g., 16 points)
    default_font_size = Pt(14)
    # Set the default font size for the entire document
    doc.styles['Normal'].font.size = default_font_size

    doc.add_paragraph(f"Login Details :                                                 {time}")
    doc.add_paragraph(f"Name: {name}                                        Email: {email}")
    doc.add_paragraph("Detected Events : ->")
    doc.add_paragraph("")

    # Add text based on flags
    if prohibited_email:
        doc.add_paragraph("Someone was detected in prohibited area")
    if fire_email:
        doc.add_paragraph("Fire was detected!!")
    if restricted_email:
        doc.add_paragraph("People count exceed limit in  restricted time")
    if accident_email:
        doc.add_paragraph("An accident was detected!!")

    # Save the modified .docx file
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    doc.save(f"reports/report_{timestamp}.docx")

